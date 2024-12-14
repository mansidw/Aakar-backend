# services/report_formatter.py
import os
from typing import List, Dict
from fpdf import FPDF
from docx import Document
from docx.shared import Inches
import base64

def image_to_base64(path: str) -> str:
    with open(path, "rb") as img_file:
        b64_data = base64.b64encode(img_file.read()).decode('utf-8')
    return f"data:image/png;base64,{b64_data}"

def generate_pdf(content: str, tables: List[Dict], charts: List[str], images: List[str]) -> str:
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    pdf.multi_cell(0, 10, content)

    for table in tables:
        pdf.ln(10)
        pdf.set_font("Arial", 'B', size=12)
        pdf.cell(0, 10, table['title'], ln=True)
        pdf.set_font("Arial", size=12)
        col_width = pdf.w / (len(table['columns']) + 1)
        # Header
        for col in table['columns']:
            pdf.cell(col_width, 10, str(col), border=1)
        pdf.ln()
        # Rows
        for row in table['rows']:
            for item in row:
                pdf.cell(col_width, 10, str(item), border=1)
            pdf.ln()

    for chart in charts:
        if os.path.exists(chart):
            pdf.ln(10)
            pdf.image(chart, x=10, w=190)

    for image in images:
        if os.path.exists(image):
            pdf.ln(10)
            pdf.image(image, x=10, w=190)

    pdf_path = "reports/report.pdf"
    os.makedirs(os.path.dirname(pdf_path), exist_ok=True)
    pdf.output(pdf_path)
    return pdf_path

def generate_docx(content: str, tables: List[Dict], charts: List[str], images: List[str]) -> str:
    document = Document()
    document.add_paragraph(content)

    for table in tables:
        document.add_heading(table['title'], level=2)
        table_obj = document.add_table(rows=1, cols=len(table['columns']))
        hdr_cells = table_obj.rows[0].cells
        for idx, col in enumerate(table['columns']):
            hdr_cells[idx].text = str(col)
        for row in table['rows']:
            row_cells = table_obj.add_row().cells
            for i, item in enumerate(row):
                row_cells[i].text = str(item)
        document.add_paragraph()

    for chart in charts:
        if os.path.exists(chart):
            document.add_heading(os.path.basename(chart), level=2)
            document.add_picture(chart, width=Inches(6))
            document.add_paragraph()

    for image in images:
        if os.path.exists(image):
            document.add_picture(image, width=Inches(6))
            document.add_paragraph()

    docx_path = "reports/report.docx"
    os.makedirs(os.path.dirname(docx_path), exist_ok=True)
    document.save(docx_path)
    return docx_path

def generate_html(content: str, tables: List[Dict], charts: List[str], images: List[str]) -> str:
    html_content = f"<html><head><title>Report</title></head><body><p>{content}</p>"

    for table in tables:
        html_content += f"<h2>{table['title']}</h2><table border='1'><tr>"
        for col in table['columns']:
            html_content += f"<th>{col}</th>"
        html_content += "</tr>"
        for row in table['rows']:
            html_content += "<tr>"
            for item in row:
                html_content += f"<td>{item}</td>"
            html_content += "</tr>"
        html_content += "</table>"

    # Convert charts to base64 for inline display
    for chart in charts:
        if os.path.exists(chart):
            b64_img = image_to_base64(chart)
            html_content += f"<h2>{os.path.basename(chart)}</h2><img src='{b64_img}' style='width:600px;'>"

    # Convert images to base64 for inline display
    for image in images:
        if os.path.exists(image):
            b64_img = image_to_base64(image)
            html_content += f"<img src='{b64_img}' style='width:600px;'>"

    html_content += "</body></html>"

    html_path = "reports/report.html"
    os.makedirs(os.path.dirname(html_path), exist_ok=True)
    with open(html_path, "w") as f:
        f.write(html_content)
    return html_path

def generate_markdown(content: str, tables: List[Dict], charts: List[str], images: List[str]) -> str:
    md_content = f"{content}\n\n"

    for table in tables:
        md_content += f"## {table['title']}\n\n"
        md_content += "| " + " | ".join(table['columns']) + " |\n"
        md_content += "|" + "|".join(["---"] * len(table['columns'])) + "|\n"
        for row in table['rows']:
            md_content += "| " + " | ".join([str(item) for item in row]) + " |\n"
        md_content += "\n"

    # Convert charts to base64
    for chart in charts:
        if os.path.exists(chart):
            b64_img = image_to_base64(chart)
            md_content += f"## {os.path.basename(chart)}\n\n"
            md_content += f"![Chart]({b64_img})\n\n"

    # Convert images to base64
    for image in images:
        if os.path.exists(image):
            b64_img = image_to_base64(image)
            md_content += f"![Image]({b64_img})\n\n"

    md_path = "reports/report.md"
    os.makedirs(os.path.dirname(md_path), exist_ok=True)
    with open(md_path, "w") as f:
        f.write(md_content)
    return md_path

def create_md_content(content: str, tables: List[Dict], charts: List[str], images: List[str]) -> str:
    md_content = f"{content}\n\n"

    for table in tables:
        md_content += f"## {table['title']}\n\n"
        md_content += "| " + " | ".join(table['columns']) + " |\n"
        md_content += "|" + "|".join(["---"] * len(table['columns'])) + "|\n"
        for row in table['rows']:
            md_content += "| " + " | ".join([str(item) for item in row]) + " |\n"
        md_content += "\n"

    for chart in charts:
        if os.path.exists(chart):
            b64_img = image_to_base64(chart)
            md_content += f"## {os.path.basename(chart)}\n\n"
            md_content += f"![Chart]({b64_img})\n\n"

    for image in images:
        if os.path.exists(image):
            b64_img = image_to_base64(image)
            md_content += f"![Image]({b64_img})\n\n"
    return md_content
