# services/file_parser.py
from llama_parse import LlamaParse
from pathlib import Path
from llama_index.llms.openai import OpenAI
from datetime import datetime
import json
from docx import Document
from openpyxl import load_workbook
import csv
import os
# Initialize LLM for generating summaries
llm = OpenAI(model="gpt-4o")

def generate_summary(text):
    """
    Generate a summary of the provided text using an LLM.
    """
    try:
        if not text:
            print("No text provided for summary generation.")
            return ""
        prompt = f"Summarize the following content: {text[:5000]}"  # Limit content for prompt length
        response = llm.complete(prompt)  # `llm.complete` returns a CompletionResponse
        summary = response.text  # Extract the text from the CompletionResponse
        print(f"Generated summary: {summary}")  # Debugging
        return summary.strip()  # Ensure the text is stripped of whitespace
    except Exception as e:
        print(f"Failed to generate summary: {e}")
        return "Summary generation failed."

def parse_file(file, file_type, image_output_dir, project_id=None):
    """
    Parse a file, generate summaries, and extract text and image chunks with metadata.
    """
    try:
        parser = LlamaParse(
            result_type="markdown",
            use_vendor_multimodal_model=True,
            vendor_multimodal_model_name="anthropic-sonnet-3.5",
            # vendor_multimodal_api_key=os.getenv("ANTHROPIC_API_KEY"),
            auto_mode=True

        )

        # Initialize text and image chunks
        text_chunks = []
        image_chunks = []
        combined_text = ""

        # Handle different file types
        if file_type == "pdf":
            md_json_objs = parser.get_json_result(file)
            print(f"Parsed JSON objects: {md_json_objs}")  # Debugging
            text_chunks = [
                {
                    "text": page.get("content", ""),  # Handle missing 'content'
                    "metadata": {
                        "page": page.get("page_num"),
                        "file_name": Path(file).name,
                        "file_type": "pdf",
                        "project_id": project_id,
                        "uploaded_at": datetime.utcnow().isoformat(),
                    }
                } 
                for page in md_json_objs[0].get("pages", [])
                if page.get("content", "").strip()  # Skip empty ch
            ]
            combined_text = " ".join(chunk["text"] for chunk in text_chunks)
            image_dicts = parser.get_images(md_json_objs, download_path=image_output_dir)
            for image in image_dicts:
                image_chunks.append({
                    "content": image.get("content", ""),  # Handle missing 'content'
                    "metadata": {
                        "page": image.get("page_num"),
                        "file_name": Path(file).name,
                        "file_type": "pdf",
                        "project_id": project_id,
                        "uploaded_at": datetime.utcnow().isoformat(),
                    }
                })

        elif file_type == "csv":
            with open(file, mode="r") as csv_file:
                csv_reader = csv.reader(csv_file)
                for row_num, row in enumerate(csv_reader, start=1):
                    row_text = " ".join(row)
                    combined_text += row_text + " "
                    text_chunks.append({
                        "text": row_text,
                        "metadata": {
                            "row": row_num,
                            "file_name": Path(file).name,
                            "file_type": "csv",
                            "project_id": project_id,
                            "uploaded_at": datetime.utcnow().isoformat(),
                        }
                    })

        elif file_type == "json":
            with open(file, mode="r") as json_file:
                data = json.load(json_file)
                combined_text = json.dumps(data, indent=2)
                text_chunks = [{
                    "text": combined_text,
                    "metadata": {
                        "file_name": Path(file).name,
                        "file_type": "json",
                        "project_id": project_id,
                        "uploaded_at": datetime.utcnow().isoformat(),
                    }
                }]

        elif file_type == "docx":
            doc = Document(file)
            combined_text = " ".join(paragraph.text for paragraph in doc.paragraphs if paragraph.text)
            text_chunks = [
                {
                    "text": paragraph.text,
                    "metadata": {
                        "file_name": Path(file).name,
                        "file_type": "docx",
                        "project_id": project_id,
                        "uploaded_at": datetime.utcnow().isoformat(),
                    }
                }
                for paragraph in doc.paragraphs if paragraph.text
            ]

        elif file_type == "xlsx":
            wb = load_workbook(file)
            for sheet in wb.worksheets:
                for row_num, row in enumerate(sheet.iter_rows(values_only=True), start=1):
                    row_text = " ".join(str(cell) for cell in row if cell is not None)
                    combined_text += row_text + " "
                    if row_text:
                        text_chunks.append({
                            "text": row_text,
                            "metadata": {
                                "sheet": sheet.title,
                                "row": row_num,
                                "file_name": Path(file).name,
                                "file_type": "xlsx",
                                "project_id": project_id,
                                "uploaded_at": datetime.utcnow().isoformat(),
                            }
                        })

        elif file_type == "txt":
            with open(file, mode="r") as txt_file:
                combined_text = txt_file.read()
                text_chunks = [{
                    "text": combined_text,
                    "metadata": {
                        "file_name": Path(file).name,
                        "file_type": "txt",
                        "project_id": project_id,
                        "uploaded_at": datetime.utcnow().isoformat(),
                    }
                }]

        else:
            return {"error": f"Unsupported file type: {file_type}"}

        # Generate a summary for the entire file
        summary = generate_summary(combined_text)

        # Add the summary to the metadata of the first text chunk
        if text_chunks:
            text_chunks[0]["metadata"]["summary"] = summary

        return {"text_chunks": text_chunks, "image_chunks": image_chunks, "summary": summary}

    except Exception as e:
        print(f"Failed to parse file: {e}")
        return None
